from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

DOC_PATH = Path('SOLUTION.docx')

content = {
    'Background': [
        'Organizations struggle with fragmented customer data, limited visibility into the sales pipeline, and inefficient lead management across departments.',
        'This CRM is designed to unify customer relationship workflows with role-based dashboards and data-driven analytics.'
    ],
    'Requirement Overview': [
        'Provide JWT-based authentication and RBAC across six role types.',
        'Support CRUD for core entities: Customers, Activities, Opportunities, Leads, Campaigns, Products, Feedback, AccountPlans, Renewals, PerformanceMetrics, DiscountRequests.',
        'Deliver role-specific dashboards and business intelligence experiences for different stakeholders.'
    ],
    'Solution Approach': [
        'Use a decoupled frontend (React + Vite + Tailwind) and backend (Node + Express + MongoDB) architecture.',
        'Implement middleware-based authorization to enforce role permissions consistently.',
        'Seed the database with realistic sample data for screenshots and analytics artifacts.'
    ],
    'Architecture Diagram': [
        'See the architecture diagram section below for a textual representation of the system layers.'
    ],
    'Technical Details': [
        'Frontend: React Router guarded routes (ProtectedRoute), context-based auth (AuthContext), and reusable components (Sidebar, Navbar, Cards).',
        'Backend: controllers and models organized by entity, JWT verification middleware, and schema validation via Mongoose.',
        'Analytics: Use seeded sample data (sample_data.xlsx) to feed dashboard metrics and Power BI visuals.'
    ],
    'Benefits': [
        'Improves sales productivity through centralized workflows and activity tracking.',
        'Enhances visibility with dashboards and performance analytics.',
        'Enables collaboration across Sales, Marketing, Product, and Executive stakeholders.'
    ],
    'Alternate Approach': [
        'Microservices split responsibilities into independent services (Auth, Customers, Sales, Activities, Reporting).',
        'This increases operational complexity compared to the current monolithic architecture for typical classroom/project scale.'
    ],
    'Assumptions': [
        'A MongoDB instance is available and accessible from the backend.',
        'JWT secret and MongoDB URI are configured via environment variables.',
        'Users will interact via a modern browser with JavaScript enabled.'
    ],
    'User Story Coverage Matrix': [
        'Sales Representative: customers, activities, opportunities, personal KPIs.',
        'Sales Manager: team performance, territory and lead assignment, forecasting.',
        'Account Manager: customer history, account plans, renewals.',
        'Marketing Team: campaigns, segmentation, lead quality and sharing.',
        'Product Manager: feedback, feature requests, roadmap planning.',
        'Executive Leadership: revenue trends, KPI cards, regional performance, win/loss analysis.'
    ],
    'Setup Instructions': [
        'Backend:\n  cd server && npm install && npm run seed && npm run dev',
        'Frontend:\n  cd client && npm install && npm run dev',
        'Login with seeded accounts (password: password123).'
    ],
    'Screenshots section': [
        'Screenshots are stored under the screenshots/ folder and are intended to demonstrate role-based navigation and dashboards.'
    ]
}

ar = [
    'Client Layer (React): Pages, dashboards, reusable components, AuthContext, API hooks',
    'API Gateway (Express): CORS, Auth middleware, routes',
    'Application Layer: controllers implement request handlers and business logic',
    'Business Logic Layer (Mongoose models): schemas and validation for each entity',
    'MongoDB Database: persistent storage and indexed collections'
]

matrix = [
    ('Sales Representative', 'Customers CRUD; Activity logging; Opportunity tracking; Personal KPIs'),
    ('Sales Manager', 'Team performance; Territory management; Lead assignment; Forecasting; Discount approval (API ready)'),
    ('Account Manager', 'Customer interaction history; Account plans; Renewal tracking'),
    ('Marketing Team', 'Campaign management; Lead generation; Segmentation; ROI tracking'),
    ('Product Manager', 'Product roadmap; Feature requests; Feedback management; Documentation'),
    ('Executive Leadership', 'Revenue analytics; KPI dashboards; Regional performance; Win/Loss and forecasting via Power BI')
]


doc = Document()

# Title
p = doc.add_paragraph('Sales Management CRM - Comprehensive Solution')
p.runs[0].font.size = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Version: 1.0')

doc.add_page_break()

for section, paragraphs in content.items():
    h = doc.add_paragraph(section)
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(14)
    for text in paragraphs:
        if '\n' in text:
            for line in text.split('\n'):
                doc.add_paragraph(line)
        else:
            doc.add_paragraph(text)

    if section == 'Architecture Diagram':
        doc.add_paragraph('System layers:')
        for line in ar:
            doc.add_paragraph(line, style='List Bullet')

    if section == 'User Story Coverage Matrix':
        doc.add_paragraph('Coverage summary:')
        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Role'
        hdr_cells[1].text = 'Covered items'
        for role, items in matrix:
            row_cells = table.add_row().cells
            row_cells[0].text = role
            row_cells[1].text = items

    doc.add_paragraph('')

DOC_PATH.unlink(missing_ok=True)
doc.save(DOC_PATH)
print('Wrote', DOC_PATH.resolve())


